from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory
from flask_mail import Mail, Message
import pymysql
from datetime import timedelta
import bcrypt
import os
import random
from werkzeug.utils import secure_filename
from config import Config

app = Flask(__name__)
app.config.from_object(Config)
app.secret_key = app.config['SECRET_KEY']
app.permanent_session_lifetime = timedelta(hours=2)
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
mail = Mail(app)

def get_db():
    connection = pymysql.connect(
        host=app.config['MYSQL_HOST'],
        user=app.config['MYSQL_USER'],
        password=app.config['MYSQL_PASSWORD'],
        database=app.config['MYSQL_DB'],
        cursorclass=pymysql.cursors.DictCursor
    )
    return connection

@app.route('/')
def index():
    return redirect(url_for('login'))

# =====================
# Login
# =====================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        try:
            con = get_db()
            cursor = con.cursor()
            cursor.execute("SELECT * FROM students WHERE email = %s", (email,))
            student = cursor.fetchone()
            con.close()
            if student and bcrypt.checkpw(password.encode('utf-8'), student['password'].encode('utf-8')):
                session['student_id'] = student['id']
                session['student_name'] = student['full_name']
                session.permanent = True
                return redirect(url_for('dashboard'))
            else:
                return render_template('index.html', error="Invalid email or password!")
        except Exception as e:
            return f"Error: {str(e)}"
    return render_template('index.html')

# =====================
# Register Step 1
# =====================
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        return render_template('register-step2.html',
            full_name   = request.form['full_name'],
            roll_no     = request.form['roll_no'],
            gender      = request.form['gender'],
            department  = request.form['department'],
            degree      = request.form['degree'],
            year        = request.form['year'],
            semester    = request.form['semester'],
            is_hosteller= request.form['is_hosteller'],
            mobile      = request.form['mobile'],
            whatsapp    = request.form['whatsapp']
        )
    return render_template('register.html')

# =====================
# Register Step 2
# =====================
@app.route('/register-step2', methods=['POST'])
def register_step2():
    email            = request.form['email']
    password         = request.form['password']
    confirm_password = request.form['confirm_password']
    full_name        = request.form['full_name']
    roll_no          = request.form['roll_no']
    gender           = request.form['gender']
    department       = request.form['department']
    degree           = request.form['degree']
    year             = request.form['year']
    semester         = request.form['semester']
    is_hosteller     = 1 if request.form['is_hosteller'] == 'Yes' else 0
    mobile           = request.form['mobile']
    whatsapp         = request.form['whatsapp']

    if password != confirm_password:
        return render_template('register-step2.html',
            error="Passwords do not match!",
            full_name=full_name, roll_no=roll_no, gender=gender,
            department=department, degree=degree, year=year,
            semester=semester, is_hosteller=request.form['is_hosteller'],
            mobile=mobile, whatsapp=whatsapp)

    hashed_pw = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())

    try:
        con = get_db()
        cursor = con.cursor()
        cursor.execute("SELECT * FROM students WHERE email = %s", (email,))
        if cursor.fetchone():
            return render_template('register-step2.html',
                error="Email already registered!",
                full_name=full_name, roll_no=roll_no, gender=gender,
                department=department, degree=degree, year=year,
                semester=semester, is_hosteller=request.form['is_hosteller'],
                mobile=mobile, whatsapp=whatsapp)

        cursor.execute("""
            INSERT INTO students 
            (full_name, roll_no, email, password, gender, department, degree, year, semester, is_hosteller, mobile, whatsapp)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (full_name, roll_no, email, hashed_pw, gender, department,
              degree, year, semester, is_hosteller, mobile, whatsapp))
        con.commit()
        con.close()
        return redirect(url_for('register_step3'))
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Register Step 3
# =====================
@app.route('/register-step3')
def register_step3():
    return render_template('register-step3.html')

# =====================
# Dashboard
# =====================
@app.route('/dashboard')
def dashboard():
    if 'student_id' not in session:
        return redirect(url_for('login'))
    try:
        con = get_db()
        cursor = con.cursor()
        sid = session['student_id']
        cursor.execute("SELECT COUNT(*) as total FROM complaints WHERE student_id = %s", (sid,))
        total = cursor.fetchone()['total']
        cursor.execute("SELECT COUNT(*) as cnt FROM complaints WHERE student_id = %s AND status = 'Pending'", (sid,))
        pending = cursor.fetchone()['cnt']
        cursor.execute("SELECT COUNT(*) as cnt FROM complaints WHERE student_id = %s AND status = 'Resolved'", (sid,))
        resolved = cursor.fetchone()['cnt']
        con.close()
        return render_template('dashboard.html',
            name=session['student_name'],
            total=total, pending=pending, resolved=resolved)
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Complaint Form
# =====================
@app.route('/complaint/new', methods=['GET', 'POST'])
def complaint_form():
    if 'student_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        category    = request.form['category']
        subject     = request.form['subject']
        priority    = request.form['priority']
        description = request.form['description']
        attachment  = None
        if 'attachment' in request.files:
            file = request.files['attachment']
            if file and file.filename != '':
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                attachment = filename
        try:
            con = get_db()
            cursor = con.cursor()
            cursor.execute("""
                INSERT INTO complaints (student_id, category, subject, priority, description, attachment, status)
                VALUES (%s, %s, %s, %s, %s, %s, 'Pending')
            """, (session['student_id'], category, subject, priority, description, attachment))
            con.commit()
            con.close()
            return redirect(url_for('my_complaints'))
        except Exception as e:
            return f"Error: {str(e)}"
    return render_template('complaint-form.html')

# =====================
# My Complaints
# =====================
@app.route('/my-complaints')
def my_complaints():
    if 'student_id' not in session:
        return redirect(url_for('login'))
    try:
        con = get_db()
        cursor = con.cursor()
        cursor.execute("SELECT * FROM complaints WHERE student_id = %s ORDER BY created_at DESC", (session['student_id'],))
        complaints = cursor.fetchall()
        con.close()
        return render_template('my-complaints.html', complaints=complaints)
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Complaint Detail
# =====================
@app.route('/complaint/<int:id>')
def complaint_detail(id):
    if 'student_id' not in session:
        return redirect(url_for('login'))
    try:
        con = get_db()
        cursor = con.cursor()
        cursor.execute("SELECT * FROM complaints WHERE id = %s AND student_id = %s", (id, session['student_id']))
        complaint = cursor.fetchone()
        con.close()
        if not complaint:
            return redirect(url_for('my_complaints'))
        return render_template('complaint-detail.html', complaint=complaint)
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Profile
# =====================
@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'student_id' not in session:
        return redirect(url_for('login'))
    try:
        con = get_db()
        cursor = con.cursor()
        if request.method == 'POST':
            cursor.execute("""
                UPDATE students SET full_name=%s, roll_no=%s, gender=%s,
                department=%s, year=%s, is_hosteller=%s, email=%s, mobile=%s, whatsapp=%s
                WHERE id=%s
            """, (
                request.form['full_name'], request.form['roll_no'],
                request.form['gender'], request.form['department'],
                request.form['year'], request.form['is_hosteller'],
                request.form['email'], request.form['mobile'],
                request.form['whatsapp'], session['student_id']
            ))
            con.commit()
            session['student_name'] = request.form['full_name']
            cursor.execute("SELECT * FROM students WHERE id = %s", (session['student_id'],))
            student = cursor.fetchone()
            con.close()
            return render_template('profile.html', student=student, success="Profile updated successfully!")
        cursor.execute("SELECT * FROM students WHERE id = %s", (session['student_id'],))
        student = cursor.fetchone()
        con.close()
        return render_template('profile.html', student=student)
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Change Password
# =====================
@app.route('/change-password', methods=['POST'])
def change_password():
    if 'student_id' not in session:
        return redirect(url_for('login'))
    try:
        old_password = request.form['old_password']
        new_password = request.form['new_password']
        con = get_db()
        cursor = con.cursor()
        cursor.execute("SELECT * FROM students WHERE id = %s", (session['student_id'],))
        student = cursor.fetchone()
        if not bcrypt.checkpw(old_password.encode('utf-8'), student['password'].encode('utf-8')):
            con.close()
            return render_template('profile.html', student=student, error="Old password is incorrect!")
        hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt())
        cursor.execute("UPDATE students SET password=%s WHERE id=%s", (hashed_pw, session['student_id']))
        con.commit()
        cursor.execute("SELECT * FROM students WHERE id = %s", (session['student_id'],))
        student = cursor.fetchone()
        con.close()
        return render_template('profile.html', student=student, success="Password changed successfully!")
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Forgot Password - Send OTP
# =====================
@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        try:
            con = get_db()
            cursor = con.cursor()
            cursor.execute("SELECT * FROM students WHERE email = %s", (email,))
            student = cursor.fetchone()

            if not student:
                con.close()
                return render_template('forgot-email.html', error="Email not found!")

            # Generate 6 digit OTP
            otp = str(random.randint(110101, 999999))

            # Delete old OTPs for this email
            cursor.execute("DELETE FROM otp_tokens WHERE email = %s", (email,))

            # Save new OTP
            cursor.execute("INSERT INTO otp_tokens (email, otp) VALUES (%s, %s)", (email, otp))
            con.commit()
            con.close()

            # Send OTP email
            msg = Message(
                subject="Your OTP - College Complaint System",
                recipients=[email]
            )
            msg.body = f"""
Hello {student['full_name']},

Your OTP for password reset is: {otp}

This OTP is valid for 10 minutes.

Do not share this OTP with anyone.

- College Complaint System
            """
            mail.send(msg)

            return render_template('forgot-otp.html', email=email, success="OTP sent to your email!")

        except Exception as e:
            return f"Error: {str(e)}"

    return render_template('forgot-email.html')

# =====================
# Verify OTP
# =====================
@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    email = request.form['email']
    otp   = request.form['otp']
    try:
        con = get_db()
        cursor = con.cursor()
        cursor.execute("SELECT * FROM otp_tokens WHERE email = %s AND otp = %s", (email, otp))
        token = cursor.fetchone()
        con.close()

        if not token:
            return render_template('forgot-otp.html', email=email, error="Invalid OTP! Try again.")

        return render_template('forgot-reset.html', email=email)

    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Reset Password
# =====================
@app.route('/reset-password', methods=['POST'])
def reset_password():
    email            = request.form['email']
    new_password     = request.form['new_password']
    confirm_password = request.form['confirm_password']

    if new_password != confirm_password:
        return render_template('forgot-reset.html', email=email, error="Passwords do not match!")

    try:
        hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt())
        con = get_db()
        cursor = con.cursor()
        cursor.execute("UPDATE students SET password=%s WHERE email=%s", (hashed_pw, email))
        cursor.execute("DELETE FROM otp_tokens WHERE email=%s", (email,))
        con.commit()
        con.close()
        return redirect(url_for('forgot_success'))
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Forgot Success
# =====================
@app.route('/forgot-success')
def forgot_success():
    return render_template('forgot-success.html')

# =====================
# Logout - Show confirmation page
# =====================
@app.route('/logout')
def logout():
    if 'student_id' not in session:
        return redirect(url_for('login'))
    return render_template('logout.html')

# =====================
# Logout Confirm - Actually logout
# =====================
@app.route('/logout-confirm')
def logout_confirm():
    session.clear()
    return redirect(url_for('login'))

# Serve Uploaded Files
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        try:
            con = get_db()
            cursor = con.cursor()
            cursor.execute("SELECT * FROM admins WHERE username = %s", (username,))
            admin = cursor.fetchone()
            con.close()
            if admin and bcrypt.checkpw(password.encode('utf-8'), admin['password'].encode('utf-8')):
                session['admin_id'] = admin['id']
                session['admin_name'] = admin['username']
                session.permanent = True
                return redirect(url_for('admin_dashboard'))
            else:
                return render_template('admin-login.html', error="Invalid username or password!")
        except Exception as e:
            return f"Error: {str(e)}"
    return render_template('admin-login.html')

# =====================
# Admin Dashboard
# =====================
@app.route('/admin/dashboard')
def admin_dashboard():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    try:
        con = get_db()
        cursor = con.cursor()

        cursor.execute("SELECT COUNT(*) as cnt FROM students")
        total_students = cursor.fetchone()['cnt']

        cursor.execute("SELECT COUNT(*) as cnt FROM complaints")
        total_complaints = cursor.fetchone()['cnt']

        cursor.execute("SELECT COUNT(*) as cnt FROM complaints WHERE status = 'Pending'")
        pending = cursor.fetchone()['cnt']

        cursor.execute("SELECT COUNT(*) as cnt FROM complaints WHERE status = 'Resolved'")
        resolved = cursor.fetchone()['cnt']

        cursor.execute("""
            SELECT c.*, s.full_name FROM complaints c
            JOIN students s ON c.student_id = s.id
            ORDER BY c.created_at DESC LIMIT 10
        """)
        recent_complaints = cursor.fetchall()
        con.close()

        return render_template('admin-dashboard.html',
            total_students=total_students,
            total_complaints=total_complaints,
            pending=pending,
            resolved=resolved,
            recent_complaints=recent_complaints
        )
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Admin All Complaints
# =====================
@app.route('/admin/complaints')
def admin_complaints():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    try:
        filter_status = request.args.get('status', '')
        search_query  = request.args.get('search', '')
        search_by     = request.args.get('search_by', 'full_name')

        # Allowed columns to prevent SQL injection
        allowed = ['full_name', 'category', 'priority', 'status']
        if search_by not in allowed:
            search_by = 'full_name'

        con = get_db()
        cursor = con.cursor()

        if search_query and filter_status:
            if search_by == 'full_name':
                cursor.execute("""
                    SELECT c.*, s.full_name FROM complaints c
                    JOIN students s ON c.student_id = s.id
                    WHERE c.status = %s AND s.full_name LIKE %s
                    ORDER BY c.created_at DESC
                """, (filter_status, f'%{search_query}%'))
            else:
                cursor.execute(f"""
                    SELECT c.*, s.full_name FROM complaints c
                    JOIN students s ON c.student_id = s.id
                    WHERE c.status = %s AND c.{search_by} LIKE %s
                    ORDER BY c.created_at DESC
                """, (filter_status, f'%{search_query}%'))
        elif search_query:
            if search_by == 'full_name':
                cursor.execute("""
                    SELECT c.*, s.full_name FROM complaints c
                    JOIN students s ON c.student_id = s.id
                    WHERE s.full_name LIKE %s
                    ORDER BY c.created_at DESC
                """, (f'%{search_query}%',))
            else:
                cursor.execute(f"""
                    SELECT c.*, s.full_name FROM complaints c
                    JOIN students s ON c.student_id = s.id
                    WHERE c.{search_by} LIKE %s
                    ORDER BY c.created_at DESC
                """, (f'%{search_query}%',))
        elif filter_status:
            cursor.execute("""
                SELECT c.*, s.full_name FROM complaints c
                JOIN students s ON c.student_id = s.id
                WHERE c.status = %s ORDER BY c.created_at DESC
            """, (filter_status,))
        else:
            cursor.execute("""
                SELECT c.*, s.full_name FROM complaints c
                JOIN students s ON c.student_id = s.id
                ORDER BY c.created_at DESC
            """)

        complaints = cursor.fetchall()
        con.close()
        return render_template('admin-complaints.html',
            complaints=complaints,
            filter_status=filter_status,
            search_query=search_query,
            search_by=search_by)
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Admin Complaint Detail
# =====================
@app.route('/admin/complaint/<int:id>')
def admin_complaint_detail(id):
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    try:
        con = get_db()
        cursor = con.cursor()
        cursor.execute("SELECT * FROM complaints WHERE id = %s", (id,))
        complaint = cursor.fetchone()
        cursor.execute("SELECT * FROM students WHERE id = %s", (complaint['student_id'],))
        student = cursor.fetchone()
        con.close()
        return render_template('admin-complaint-detail.html', complaint=complaint, student=student)
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Admin Update Status
# =====================
@app.route('/admin/update-status', methods=['POST'])
def admin_update_status():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    try:
        complaint_id = request.form['complaint_id']
        status = request.form['status']
        con = get_db()
        cursor = con.cursor()
        cursor.execute("UPDATE complaints SET status = %s WHERE id = %s", (status, complaint_id))
        con.commit()
        cursor.execute("SELECT * FROM complaints WHERE id = %s", (complaint_id,))
        complaint = cursor.fetchone()
        cursor.execute("SELECT * FROM students WHERE id = %s", (complaint['student_id'],))
        student = cursor.fetchone()
        con.close()

        # Send Email Notification to Student
        complaint_id_str = 'CMP' + str(complaint['id']).zfill(3)

        if status == 'Under Review':
            subject = f"Your Complaint {complaint_id_str} is Under Review"
            body = f"""Dear {student['full_name']},

Your complaint has been received and is currently under review by our team.

Complaint Details:
- Complaint ID  : {complaint_id_str}
- Category      : {complaint['category']}
- Subject       : {complaint['subject']}
- Status        : Under Review

We will update you once the issue is resolved.

Regards,
College Complaint System"""

        elif status == 'Resolved':
            subject = f"Your Complaint {complaint_id_str} has been Resolved!"
            body = f"""Dear {student['full_name']},

Great news! Your complaint has been resolved successfully.

Complaint Details:
- Complaint ID  : {complaint_id_str}
- Category      : {complaint['category']}
- Subject       : {complaint['subject']}
- Status        : Resolved

Thank you for bringing this to our attention.

Regards,
College Complaint System"""

        elif status == 'Rejected':
            subject = f"Update on Your Complaint {complaint_id_str}"
            body = f"""Dear {student['full_name']},

We regret to inform you that your complaint has been reviewed and rejected.

Complaint Details:
- Complaint ID  : {complaint_id_str}
- Category      : {complaint['category']}
- Subject       : {complaint['subject']}
- Status        : Rejected

If you believe this is incorrect, please submit a new complaint with more details.

Regards,
College Complaint System"""

        else:
            subject = f"Update on Your Complaint {complaint_id_str}"
            body = f"""Dear {student['full_name']},

Your complaint status has been updated.

Complaint Details:
- Complaint ID  : {complaint_id_str}
- Category      : {complaint['category']}
- Subject       : {complaint['subject']}
- Status        : {status}

Regards,
College Complaint System"""

        msg = Message(subject=subject, recipients=[student['email']])
        msg.body = body
        mail.send(msg)

        return render_template('admin-complaint-detail.html',
            complaint=complaint, student=student,
            success=f"Status updated to '{status}' and student notified via email! 📧")
    except Exception as e:
        return f"Error: {str(e)}"

# =====================
# Admin All Students
# =====================
@app.route('/admin/students')
def admin_students():
    if 'admin_id' not in session:
        return redirect(url_for('admin_login'))
    try:
        search_query = request.args.get('search', '')
        search_by    = request.args.get('search_by', 'full_name')

        # Allowed columns to prevent SQL injection
        allowed = ['full_name', 'roll_no', 'email', 'department', 'year', 'mobile']
        if search_by not in allowed:
            search_by = 'full_name'

        con = get_db()
        cursor = con.cursor()

        if search_query:
            cursor.execute(f"""
                SELECT * FROM students
                WHERE {search_by} LIKE %s
                ORDER BY created_at DESC
            """, (f'%{search_query}%',))
        else:
            cursor.execute("SELECT * FROM students ORDER BY created_at DESC")

        students = cursor.fetchall()
        con.close()
        return render_template('admin-students.html',
            students=students,
            search_query=search_query,
            search_by=search_by)
    except Exception as e:
        return f"Error: {str(e)}"
@app.route('/admin/logout')
def admin_logout():
    session.pop('admin_id', None)
    session.pop('admin_name', None)
    return redirect(url_for('admin_login'))


# =====================
# TUTOR SECTION
# =====================

# Tutor Login
@app.route('/tutor/login', methods=['GET', 'POST'])
def tutor_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        try:
            con = get_db()
            cursor = con.cursor()
            cursor.execute("SELECT * FROM tutors WHERE username = %s", (username,))
            tutor = cursor.fetchone()
            con.close()
            if tutor and tutor['password'] == password:
                session['tutor_id'] = tutor['id']
                session['tutor_name'] = tutor['name']
                session.permanent = True
                return redirect(url_for('tutor_dashboard'))
            else:
                return render_template('tutor-login.html', error="Invalid username or password!")
        except Exception as e:
            return f"Error: {str(e)}"
    return render_template('tutor-login.html')

# Tutor Dashboard
@app.route('/tutor/dashboard')
def tutor_dashboard():
    if 'tutor_id' not in session:
        return redirect(url_for('tutor_login'))
    try:
        con = get_db()
        cursor = con.cursor()
        tutor_id = session['tutor_id']

        # Stats
        cursor.execute("SELECT COUNT(DISTINCT session_id) FROM ccm_sessions WHERE tutor_id = %s", (tutor_id,))
        total_sessions = cursor.fetchone()['COUNT(DISTINCT session_id)'] or 0

        cursor.execute("SELECT COUNT(DISTINCT session_id) FROM ccm_sessions WHERE tutor_id = %s AND status = 'active'", (tutor_id,))
        active_sessions = cursor.fetchone()['COUNT(DISTINCT session_id)'] or 0

        cursor.execute("""
            SELECT COUNT(*) as cnt FROM ccm_responses r
            JOIN ccm_sessions s ON r.session_id = s.session_id
            WHERE s.tutor_id = %s
        """, (tutor_id,))
        total_responses = cursor.fetchone()['cnt'] or 0

        cursor.execute("""
            SELECT COUNT(*) as cnt FROM ccm_sessions s
            WHERE s.tutor_id = %s AND s.status = 'active'
            AND s.roll_number NOT IN (SELECT roll_number FROM ccm_responses WHERE session_id = s.session_id)
        """, (tutor_id,))
        pending_responses = cursor.fetchone()['cnt'] or 0

        # Recent sessions
        cursor.execute("""
            SELECT s.session_id as id, s.subject_code, s.subject_name, s.professor_name,
                   s.status, s.created_at,
                   COUNT(DISTINCT s2.roll_number) as member_count,
                   COUNT(DISTINCT r.roll_number) as response_count
            FROM ccm_sessions s
            LEFT JOIN ccm_sessions s2 ON s.session_id = s2.session_id
            LEFT JOIN ccm_responses r ON s.session_id = r.session_id
            WHERE s.tutor_id = %s
            GROUP BY s.session_id, s.subject_code, s.subject_name, s.professor_name, s.status, s.created_at
            ORDER BY s.created_at DESC LIMIT 10
        """, (tutor_id,))
        recent_sessions = cursor.fetchall()
        con.close()

        tutor = {'name': session['tutor_name']}
        return render_template('tutor-dashboard.html',
            tutor=tutor,
            total_sessions=total_sessions,
            active_sessions=active_sessions,
            total_responses=total_responses,
            pending_responses=pending_responses,
            recent_sessions=recent_sessions)
    except Exception as e:
        return f"Error: {str(e)}"

# Tutor CCM Page
@app.route('/tutor/ccm')
def tutor_ccm():
    if 'tutor_id' not in session:
        return redirect(url_for('tutor_login'))
    try:
        con = get_db()
        cursor = con.cursor()
        tutor_id = session['tutor_id']

        # Check active session
        cursor.execute("""
            SELECT s.*, 
                   (SELECT COUNT(*) FROM ccm_responses r WHERE r.session_id = s.session_id AND r.roll_number = s.boy_roll) as boy_responded,
                   (SELECT COUNT(*) FROM ccm_responses r WHERE r.session_id = s.session_id AND r.roll_number = s.girl_roll) as girl_responded
            FROM ccm_sessions s
            WHERE s.tutor_id = %s AND s.status = 'active'
            ORDER BY s.created_at DESC
        """, (tutor_id,))
        active_session = cursor.fetchall()

        # Past sessions grouped by session_id
        cursor.execute("""
            SELECT s.session_id, s.subject_code, s.subject_name, s.professor_name, s.created_at,
                   COUNT(DISTINCT s.roll_number) as member_count,
                   COUNT(DISTINCT r.roll_number) as response_count
            FROM ccm_sessions s
            LEFT JOIN ccm_responses r ON s.session_id = r.session_id
            WHERE s.tutor_id = %s AND s.status = 'stopped'
            GROUP BY s.session_id, s.subject_code, s.subject_name, s.professor_name, s.created_at
            ORDER BY s.created_at DESC
        """, (tutor_id,))
        past_sessions = cursor.fetchall()
        con.close()

        return render_template('tutor-ccm.html',
            active_session=active_session if active_session else None,
            past_sessions=past_sessions)
    except Exception as e:
        return f"Error: {str(e)}"

# Start CCM Session
@app.route('/tutor/ccm/start', methods=['POST'])
def tutor_ccm_start():
    if 'tutor_id' not in session:
        return redirect(url_for('tutor_login'))
    try:
        tutor_id = session['tutor_id']
        subject_codes   = request.form.getlist('subject_code[]')
        subject_names   = request.form.getlist('subject_name[]')
        professor_names = request.form.getlist('professor_name[]')
        boy_rolls       = request.form.getlist('boy_roll[]')
        girl_rolls      = request.form.getlist('girl_roll[]')

        import uuid
        session_id = str(uuid.uuid4())[:8].upper()

        con = get_db()
        cursor = con.cursor()

        for i in range(len(subject_codes)):
            # Insert boy member
            cursor.execute("""
                INSERT INTO ccm_sessions
                (session_id, tutor_id, subject_code, subject_name, professor_name, roll_number, gender, boy_roll, girl_roll, status)
                VALUES (%s, %s, %s, %s, %s, %s, 'boy', %s, %s, 'active')
            """, (session_id, tutor_id, subject_codes[i], subject_names[i],
                  professor_names[i], boy_rolls[i], boy_rolls[i], girl_rolls[i]))
            # Insert girl member
            cursor.execute("""
                INSERT INTO ccm_sessions
                (session_id, tutor_id, subject_code, subject_name, professor_name, roll_number, gender, boy_roll, girl_roll, status)
                VALUES (%s, %s, %s, %s, %s, %s, 'girl', %s, %s, 'active')
            """, (session_id, tutor_id, subject_codes[i], subject_names[i],
                  professor_names[i], girl_rolls[i], boy_rolls[i], girl_rolls[i]))

        con.commit()
        con.close()
        return redirect(url_for('tutor_ccm'))
    except Exception as e:
        return f"Error: {str(e)}"

# Stop CCM Session
@app.route('/tutor/ccm/stop', methods=['POST'])
def tutor_ccm_stop():
    if 'tutor_id' not in session:
        return redirect(url_for('tutor_login'))
    try:
        tutor_id = session['tutor_id']
        con = get_db()
        cursor = con.cursor()
        cursor.execute("UPDATE ccm_sessions SET status = 'stopped' WHERE tutor_id = %s AND status = 'active'", (tutor_id,))
        con.commit()
        con.close()
        return redirect(url_for('tutor_ccm'))
    except Exception as e:
        return f"Error: {str(e)}"

# Download CCM Responses as Word
@app.route('/tutor/ccm/download')
@app.route('/tutor/ccm/download/<session_id>')
def tutor_ccm_download(session_id=None):
    if 'tutor_id' not in session:
        return redirect(url_for('tutor_login'))
    try:
        from docx import Document
        from io import BytesIO
        from flask import send_file
        import datetime

        tutor_id = session['tutor_id']
        con = get_db()
        cursor = con.cursor()

        if session_id:
            cursor.execute("""
                SELECT s.subject_code, s.subject_name, s.professor_name,
                       r.roll_number, r.grievance, r.submitted_at
                FROM ccm_sessions s
                LEFT JOIN ccm_responses r ON s.session_id = r.session_id AND s.roll_number = r.roll_number
                WHERE s.session_id = %s AND s.tutor_id = %s
                ORDER BY s.subject_code, s.roll_number
            """, (session_id, tutor_id))
        else:
            cursor.execute("""
                SELECT s.subject_code, s.subject_name, s.professor_name,
                       r.roll_number, r.grievance, r.submitted_at
                FROM ccm_sessions s
                LEFT JOIN ccm_responses r ON s.session_id = r.session_id AND s.roll_number = r.roll_number
                WHERE s.tutor_id = %s AND s.status = 'active'
                ORDER BY s.subject_code, s.roll_number
            """, (tutor_id,))

        rows = cursor.fetchall()
        con.close()

        # Build Word Document
        doc = Document()
        doc.add_heading('CCM Grievance Report', 0)
        doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%d-%m-%Y %H:%M")}')
        doc.add_paragraph(f'Tutor: {session["tutor_name"]}')
        doc.add_paragraph('')

        # Group by subject
        subjects = {}
        for row in rows:
            key = row['subject_code']
            if key not in subjects:
                subjects[key] = []
            subjects[key].append(row)

        for subj_code, entries in subjects.items():
            doc.add_heading(f'{subj_code} — {entries[0]["subject_name"]}', level=1)
            doc.add_paragraph(f'Professor: {entries[0]["professor_name"]}')

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Roll Number'
            hdr[1].text = 'Grievance'
            hdr[2].text = 'Submitted At'

            for entry in entries:
                row_cells = table.add_row().cells
                row_cells[0].text = entry['roll_number'] or '-'
                row_cells[1].text = entry['grievance'] or 'No response yet'
                row_cells[2].text = entry['submitted_at'].strftime('%d-%m-%Y %H:%M') if entry['submitted_at'] else '-'

            doc.add_paragraph('')

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f'CCM_Report_{datetime.datetime.now().strftime("%d%m%Y_%H%M")}.docx'
        return send_file(buf, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return f"Error: {str(e)}"

# Tutor Logout
@app.route('/tutor/logout')
def tutor_logout():
    session.pop('tutor_id', None)
    session.pop('tutor_name', None)
    return redirect(url_for('tutor_login'))


if __name__ == '__main__':
    app.run(debug=True)
